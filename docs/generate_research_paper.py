"""Generate SafePay IEEE-style research paper DOCX with figures and optional ML metrics."""

from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

DOCS = Path(__file__).resolve().parent
ASSETS = DOCS / "paper_assets"
OUTPUT = DOCS / "SafePay Research paper.docx"
OUTPUT_FALLBACK = DOCS / "SafePay Research paper (draft).docx"
METRICS_JSON = ASSETS / "ml_evaluation_results.json"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_figure(doc, image_path: Path, caption: str, width=Inches(5.8)):
    if image_path.exists():
        doc.add_picture(str(image_path), width=width)
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.italic = True
            run.font.size = Pt(10)
    else:
        add_para(doc, f"[Figure missing: {image_path.name}]", italic=True)


def add_table(doc, headers, rows, title=None):
    if title:
        add_heading(doc, title, 2)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = str(val)


def fmt_metric(m: dict | None, key: str, as_pct=True) -> str:
    if not m or key not in m:
        return "TBD"
    val = m[key]
    if as_pct and key not in ("tp", "tn", "fp", "fn"):
        return f"{val * 100:.2f}%"
    return str(val)


def load_metrics():
    if not METRICS_JSON.exists():
        return None
    return json.loads(METRICS_JSON.read_text(encoding="utf-8"))


def model_row(metrics_data, name: str):
    models = {m["model"]: m for m in metrics_data.get("models", [])} if metrics_data else {}
    m = models.get(name)
    return [
        name,
        fmt_metric(m, "accuracy"),
        fmt_metric(m, "precision"),
        fmt_metric(m, "recall"),
        fmt_metric(m, "f1"),
        fmt_metric(m, "roc_auc"),
    ]


def generate_figures():
    script = ASSETS / "generate_figures.py"
    if script.exists():
        subprocess.run([sys.executable, str(script)], check=True, cwd=str(ASSETS))


def build_abstract(metrics_data) -> str:
    xgb = None
    if metrics_data:
        xgb = next((m for m in metrics_data["models"] if m["model"] == "XGBoost"), None)

    base = (
        "Digital payment fraud is increasing globally, yet many consumer payment systems "
        "still rely on static identity checks and post-transaction rule engines. These "
        "approaches are reactive, institution-siloed, and difficult to audit when an "
        "automated system blocks a legitimate user. This paper presents SafePay, a "
        "multi-layer fraud-detection framework that combines behavioural biometrics, "
        "gradient-boosted machine learning, explainable AI, blockchain-based anonymised "
        "fraud-intelligence sharing, and federated learning within a real-time payment "
        "pipeline. SafePay was implemented as a containerised microservices platform "
        "(Next.js, FastAPI, PostgreSQL, Redis, XGBoost, Solidity/Hardhat, Flower) and "
        "evaluated on the IEEE-CIS Fraud Detection benchmark (590,540 transactions). "
    )

    if xgb:
        ml_part = (
            f"The centralised XGBoost classifier achieved {fmt_metric(xgb, 'accuracy')} test accuracy, "
            f"{fmt_metric(xgb, 'precision')} precision, {fmt_metric(xgb, 'recall')} recall, "
            f"{fmt_metric(xgb, 'f1')} F1-score, and {fmt_metric(xgb, 'roc_auc')} ROC-AUC "
            f"after preprocessing 394 raw features to 338 usable features. "
        )
    else:
        ml_part = (
            "The centralised XGBoost classifier achieved 97.55% test accuracy on a hold-out set "
            "(full precision/recall/F1/ROC-AUC pending baseline evaluation run). "
        )

    tail = (
        "Federated training across three simulated bank clients using Flower FedXgbBagging "
        "reached validation AUC 0.846 without exchanging raw transaction data. SHAP was "
        "integrated at inference time for challenged/blocked transactions. System components "
        "were validated through live API testing. Limitations include benchmark-only training "
        "data, simulated federation, local blockchain testnet, and incomplete ablation "
        "experiments for behavioural/device layers."
    )
    return base + ml_part + tail


def build_document():
    generate_figures()
    metrics_data = load_metrics()
    xgb = None
    if metrics_data:
        xgb = next((m for m in metrics_data["models"] if m["model"] == "XGBoost"), None)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(
        "SafePay: A Privacy-Preserving Multi-Layer AI Framework "
        "for Real-Time Financial Fraud Detection"
    )
    tr.bold = True
    tr.font.size = Pt(14)

    for line in [
        "Jay Jadhav, Ganesh, Mayur",
        "Department of Information Technology",
        "Prof. Ram Meghe Institute of Technology & Research, Badnera",
        "Amravati, India — Academic Year 2025–2026",
        "Guide: Prof. Pravin Nerkar",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_paragraph()
    add_heading(doc, "Abstract", 1)
    add_para(doc, build_abstract(metrics_data))

    add_heading(doc, "Index Terms", 1)
    add_para(
        doc,
        "Financial fraud detection, machine learning, XGBoost, behavioural biometrics, "
        "federated learning, blockchain, explainable AI, digital payment security.",
    )

    # --- I. INTRODUCTION ---
    add_heading(doc, "I. INTRODUCTION", 1)
    add_para(
        doc,
        "India's Unified Payments Interface (UPI) and digital-wallet ecosystems have "
        "dramatically increased payment convenience, but they have also expanded the attack "
        "surface for credential theft, social engineering, and device-farm fraud. Traditional "
        "systems verify identity at login using OTP, PIN, or biometrics, yet they rarely "
        "verify whether ongoing session behaviour matches the legitimate account owner.",
    )
    add_para(
        doc,
        "Existing fraud platforms exhibit four recurring limitations: (1) reactive detection "
        "that flags suspicious activity after settlement; (2) data silos that prevent one "
        "institution from learning from another's confirmed fraud; (3) static Know Your "
        "Customer (KYC) checks that do not continuously authenticate users; and (4) black-box "
        "model decisions that are difficult to explain to users, analysts, or regulators.",
    )
    add_para(doc, "The main contributions of this work are:")
    add_bullets(
        doc,
        [
            "A multi-layer real-time fraud-detection architecture integrating behavioural, "
            "transaction, device, and ML signals in a single weighted scoring pipeline.",
            "Behavioural telemetry collection (keystroke, touch, device fingerprint) linked "
            "to payment risk assessment.",
            "Explainable fraud decisions using SHAP TreeExplainer on a 338-feature XGBoost model.",
            "Privacy-preserving cross-institution fraud-signal sharing using keccak256-hashed "
            "identifiers on Ethereum smart contracts — no raw PII on chain.",
            "Federated learning simulation with Flower across three bank clients (FedXgbBagging).",
            "Experimental evaluation of ML classifiers on the IEEE-CIS benchmark with baseline "
            "comparison, plus system-level validation of authentication, payments, and security controls.",
        ],
    )

    # --- II. RELATED WORK ---
    add_heading(doc, "II. RELATED WORK", 1)
    add_para(
        doc,
        "Recent fraud-detection research on IEEE-CIS shows that gradient-boosted tree models "
        "outperform linear baselines on tabular transaction data, but most published systems "
        "score transactions in batch mode and do not integrate continuous behavioural biometrics "
        "into a live payment path. Behavioural biometrics surveys report strong user "
        "identification in controlled sessions, yet few works connect keystroke/touch telemetry "
        "to real-time wallet authorisation. Federated learning literature (FedAvg and variants) "
        "demonstrates near-centralised performance while keeping data local. Blockchain fraud "
        "proposals emphasise hash-based cross-institution signals rather than plaintext PII. "
        "SHAP and related XAI methods are increasingly cited for financial decision auditability.",
    )

    add_table(
        doc,
        ["Reference", "Year", "Method", "Dataset", "Result", "Limitation"],
        [
            ["[6] Abdallah et al.", "2026", "Ensemble + SMOTE", "IEEE-CIS", "High AUC", "Centralised; no biometrics"],
            ["[11] Fiore et al.", "2019", "Gradient boosting", "IEEE-CIS", "Strong AUC", "Offline scoring only"],
            ["[7] Mondal & Bours", "2023", "Keystroke dynamics", "Lab users", ">95% EER gains", "Not payment-integrated"],
            ["[1] McMahan et al.", "2017", "FedAvg", "Decentralised shards", "Near-central accuracy", "No blockchain / XAI"],
            ["[8] Zheng et al.", "2023", "Blockchain fraud ledger", "Simulated banks", "Tamper-proof sharing", "No real-time ML pipeline"],
            ["[2] Lundberg & Lee", "2017", "SHAP", "General ML", "Unified attributions", "Not a fraud system alone"],
            ["This work (SafePay)", "2026", "Multi-layer framework", "IEEE-CIS + live app", "See Tables II–IV", "Benchmark + testnet scope"],
        ],
        title="TABLE I — COMPARATIVE LITERATURE SUMMARY",
    )

    # --- III. RESEARCH GAP ---
    add_heading(doc, "III. RESEARCH GAP", 1)
    add_para(
        doc,
        "Rather than claiming that no prior system exists, SafePay targets a specific integration "
        "gap: we did not identify a widely cited single architecture that simultaneously provides "
        "(a) real-time payment-path scoring, (b) continuous behavioural biometrics, (c) explainable "
        "ML output, (d) anonymised cross-institution signal lookup, and (e) federated model training "
        "within one deployable platform. SafePay implements this integration; quantitative evidence "
        "for each layer is summarised in Table IV.",
    )

    # --- IV. ARCHITECTURE ---
    add_heading(doc, "IV. PROPOSED SYSTEM ARCHITECTURE", 1)
    add_para(
        doc,
        "SafePay follows a microservices architecture with six Docker containers: frontend "
        "(Next.js 14), backend (FastAPI), ml-service (XGBoost + SHAP), PostgreSQL, Redis, "
        "and a local Hardhat Ethereum node. The fraud engine executes before payment commit.",
    )
    add_para(
        doc,
        "Weighted risk formula (fraud_service.py): "
        "Score = 0.35×BehavioralRisk + 0.30×TransactionRisk + 0.20×DeviceRisk + 0.15×MLScore. "
        "Thresholds: approve if score < 0.30; challenge if 0.30–0.70; block if > 0.70.",
    )

    add_figure(
        doc,
        ASSETS / "fig1_architecture.png",
        "Fig. 1. SafePay real-time fraud detection pipeline (main payment path).",
    )
    add_figure(
        doc,
        ASSETS / "fig3_supporting_layers.png",
        "Fig. 3. Supporting layers: blockchain fraud intelligence and federated learning.",
        width=Inches(5.5),
    )

    add_table(
        doc,
        ["Layer", "Implementation", "Evaluation"],
        [
            ["Transaction ML", "XGBoost (338 features)", "Table II: Precision/Recall/F1/AUC"],
            ["Behavioural", "Telemetry + heuristic trust score", "Scenario tests; ablation TBD (Table V)"],
            ["Device", "Fingerprint + trust score", "Scenario tests; ablation TBD (Table V)"],
            ["SHAP / XAI", "TreeExplainer, top-5 factors", "Explanation examples in SOC UI"],
            ["Blockchain", "FraudRegistry + Reputation (Hardhat)", "Lookup + publish on confirmed fraud"],
            ["Federated", "Flower, 3 clients, FedXgbBagging", "Table IV: FL vs centralised AUC"],
        ],
        title="TABLE IV — LAYER IMPLEMENTATION AND EVALUATION STATUS",
    )

    # --- V. COMPONENTS ---
    add_heading(doc, "V. SYSTEM COMPONENTS", 1)
    components = [
        ("Authentication", "Registration, OTP (Redis HMAC-SHA256), JWT rotation, logout revocation, RBAC (5 roles)."),
        ("Payment engine", "Wallet, P2P, merchant pay, QR, UPI-style send, idempotency, row-level wallet locking."),
        ("Behavioural biometrics", "Device fingerprint middleware, /behavior/telemetry, trust score 0–100, baseline ≥20 events."),
        ("Fraud ML", "XGBoost on IEEE-CIS; /score returns risk_score, decision, SHAP contributions."),
        ("Blockchain", "keccak256(entity+salt); auto-publish on confirmed fraud; lookup/reputation APIs."),
        ("Federated learning", "3 simulated banks; disjoint shards (~196,846 rows each via modulo partition); FedXgbBagging."),
        ("AI copilot", "LangGraph + Gemini 1.5 Flash; grounded transaction/risk explanations."),
    ]
    for name, desc in components:
        add_para(doc, f"{name}: {desc}", bold=True)

    # --- VI. METHODOLOGY ---
    add_heading(doc, "VI. METHODOLOGY", 1)

    add_heading(doc, "A. Dataset and Preprocessing", 2)
    add_para(
        doc,
        "Dataset: IEEE-CIS Fraud Detection (Kaggle, 2019). Records: 590,540. Raw features: 394. "
        "Target: isFraud (~3.5% fraud). Steps: drop columns with >80% missing (394→339), fill NaN "
        "with 0, LabelEncoder on categoricals, 80/20 stratified split (random_state=42), 338 features.",
    )

    add_heading(doc, "B. Data Leakage Controls", 2)
    leakage = metrics_data.get("data_leakage_controls") if metrics_data else [
        "LabelEncoders must be fit on training split only",
        "Any imputation/scaling fit on training split only",
        "No test labels used in feature selection",
        "SMOTE (if used) must be applied only to training data",
    ]
    add_bullets(doc, leakage)

    add_figure(
        doc,
        ASSETS / "fig2_experimental_flow.png",
        "Fig. 2. Offline ML evaluation workflow with train-only preprocessing.",
    )

    add_heading(doc, "C. XGBoost Classifier", 2)
    add_para(
        doc,
        "Hyperparameters from train_model.py (not tuned via grid search): n_estimators=50, "
        "max_depth=6, learning_rate=0.1, eval_metric=logloss, random_state=42. These values "
        "were chosen as a reproducible baseline consistent with the team's training script.",
    )

    add_heading(doc, "D. Baseline Models", 2)
    add_para(
        doc,
        "For comparative evaluation we train Logistic Regression (balanced class weights + "
        "StandardScaler), Decision Tree (max_depth=12, balanced), Random Forest (100 trees), "
        "and XGBoost under identical train/test splits. Metrics: Accuracy, Precision, Recall, "
        "F1, ROC-AUC, PR-AUC, FPR, FNR (see evaluate_models.py).",
    )

    add_heading(doc, "E. SHAP Explanation", 2)
    add_para(
        doc,
        "TreeExplainer computes SHAP values for the fraud class; top-5 |SHAP| features are "
        "returned per prediction and stored in fraud_explanations.",
    )

    add_heading(doc, "F. Federated Learning Protocol", 2)
    add_para(
        doc,
        "Three simulated bank clients (bank_a, bank_b, bank_c) train on disjoint partitions "
        "of IEEE-CIS (round-robin index modulo 3). Each client performs an 80/20 local "
        "train/validation split (stratified). Flower FedXgbBagging aggregates XGBoost updates. "
        "Local params: eta=0.1, max_depth=8, eval_metric=auc. Reported federated validation "
        "AUC: 0.846. Only aggregated metrics — not raw rows — leave client processes.",
    )

    # --- VII. EXPERIMENTAL SETUP ---
    add_heading(doc, "VII. EXPERIMENTAL SETUP", 1)
    add_para(
        doc,
        "Software: Python 3.11+, FastAPI, XGBoost 2.x, scikit-learn, SHAP 0.44+, Flower, "
        "Solidity 0.8.24, Hardhat, Docker Compose. Hardware: development laptop (CPU). "
        "Offline ML evaluation uses the script ml-service/app/models/evaluate_models.py. "
        "Online validation uses live curl/API tests against Docker containers.",
    )

    add_table(
        doc,
        ["Pipeline Stage", "Mean (ms)", "P95 (ms)", "P99 (ms)", "Notes"],
        [
            ["Authentication", "TBD", "TBD", "TBD", "<100 ms observed informally — formal benchmark pending"],
            ["Payment (excl. fraud)", "TBD", "TBD", "TBD", "<150 ms observed informally"],
            ["ML /score only", "TBD", "TBD", "TBD", "Target <500 ms — measured benchmark pending"],
            ["Complete fraud pipeline", "TBD", "TBD", "TBD", "Do not claim sub-second until measured"],
        ],
        title="TABLE VI — LATENCY BENCHMARK (TO MEASURE BEFORE SUBMISSION)",
    )

    # --- VIII. RESULTS ---
    add_heading(doc, "VIII. RESULTS", 1)

    add_table(
        doc,
        ["Model", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC"],
        [
            model_row(metrics_data, "Logistic Regression"),
            model_row(metrics_data, "Decision Tree"),
            model_row(metrics_data, "Random Forest"),
            model_row(metrics_data, "XGBoost"),
        ],
        title="TABLE II — MODEL COMPARISON ON IEEE-CIS TEST SET",
    )

    if not metrics_data:
        add_para(
            doc,
            "NOTE: Run `py -m app.models.evaluate_models --data data/train_transaction.csv` "
            "from ml-service/ to populate Table II and Figs. 4–6. Do not estimate metrics from accuracy alone.",
            italic=True,
        )
    elif xgb:
        add_para(
            doc,
            f"XGBoost PR-AUC: {fmt_metric(xgb, 'pr_auc')}. "
            f"FPR: {fmt_metric(xgb, 'fpr')}. FNR: {fmt_metric(xgb, 'fnr')}. "
            "On imbalanced fraud data, recall and FNR are critical: false negatives are "
            "fraudulent transactions classified as genuine; false positives block legitimate users.",
        )

    add_table(
        doc,
        ["Actual \\ Predicted", "Genuine", "Fraud"],
        [
            [
                "Genuine",
                f"TN = {xgb['tn'] if xgb else 'TBD'}",
                f"FP = {xgb['fp'] if xgb else 'TBD'}",
            ],
            [
                "Fraud",
                f"FN = {xgb['fn'] if xgb else 'TBD'}",
                f"TP = {xgb['tp'] if xgb else 'TBD'}",
            ],
        ],
        title="TABLE III — CONFUSION MATRIX (XGBOOST, TEST SET)",
    )

    add_figure(
        doc,
        ASSETS / "fig4_confusion_matrix.png",
        "Fig. 4. XGBoost confusion matrix on IEEE-CIS test set.",
        width=Inches(4.2),
    )
    add_figure(
        doc,
        ASSETS / "fig5_roc_curves.png",
        "Fig. 5. ROC curves — baseline model comparison.",
        width=Inches(4.8),
    )
    add_figure(
        doc,
        ASSETS / "fig6_pr_curves.png",
        "Fig. 6. Precision-recall curves — baseline model comparison.",
        width=Inches(4.8),
    )

    add_table(
        doc,
        ["Model", "ROC-AUC", "Notes"],
        [
            ["Centralised XGBoost", fmt_metric(xgb, "roc_auc") if xgb else "TBD", "Same test split as Table II"],
            ["Federated XGBoost (FedXgbBagging)", "0.846", "3-client Flower simulation; validation AUC"],
            ["Centralised vs FL gap", "TBD", "Report absolute AUC difference after centralised AUC is finalised"],
        ],
        title="TABLE VII — CENTRALISED VS FEDERATED AUC COMPARISON",
    )

    add_table(
        doc,
        ["Configuration", "Accuracy", "Precision", "Recall", "F1", "AUC"],
        [
            ["Transaction features only", "TBD", "TBD", "TBD", "TBD", "TBD"],
            ["+ Device signals", "TBD", "TBD", "TBD", "TBD", "TBD"],
            ["+ Behavioural signals", "TBD", "TBD", "TBD", "TBD", "TBD"],
            ["+ ML score (full SafePay pipeline)", "TBD", "TBD", "TBD", "TBD", "TBD"],
        ],
        title="TABLE V — ABLATION STUDY (TO COMPLETE — REQUIRES CONTROLLED SCENARIO RUNS)",
    )

    add_para(
        doc,
        "Security validation (live API): refresh-token reuse → 401; logout revocation verified; "
        "idempotency prevents double-spend; wallet race condition prevented via SELECT ... FOR UPDATE.",
    )

    # --- IX. DISCUSSION ---
    add_heading(doc, "IX. DISCUSSION", 1)
    if metrics_data:
        add_para(
            doc,
            "Baseline comparison (Table II) allows stating whether XGBoost outperforms linear and "
            "tree baselines on F1 and ROC-AUC — only if the measured results support that claim.",
        )
    add_para(
        doc,
        "The framework is designed to incorporate behavioural and device signals alongside "
        "transaction-level ML in the composite score. Controlled ablation experiments (Table V) "
        "are required before claiming that behavioural/device layers improve recall over "
        "transaction-only scoring. SHAP bridges model output and human trust for challenge/block "
        "decisions. Blockchain signals use hashed identifiers; immutability supports cross-institution "
        "auditability though PostgreSQL could store signals in a centralised deployment.",
    )
    add_para(
        doc,
        "False negatives (FN) represent missed fraud and are typically more costly than false "
        "positives (FP) in security-sensitive deployments, though excessive FP harms user experience. "
        "The final paper should discuss this trade-off using measured FN/FP from Table III.",
    )

    # --- X. LIMITATIONS ---
    add_heading(doc, "X. LIMITATIONS", 1)
    add_bullets(
        doc,
        [
            "Training uses the public IEEE-CIS benchmark, not proprietary banking data.",
            "Behavioural biometrics use heuristic trust scoring, not a dedicated retrained classifier.",
            "Blockchain runs on a local Hardhat testnet, not production mainnet.",
            "Federated learning uses simulated bank clients on partitioned benchmark shards.",
            "No real UPI/NPCI integration; payments use SafePay sandbox rails.",
            "Baseline/ablation tables and P95/P99 latency benchmarks incomplete until experiments are run.",
            "Adversarial attacks against behavioural or ML components not fully evaluated.",
        ],
    )

    # --- XI. FUTURE WORK ---
    add_heading(doc, "XI. FUTURE WORK", 1)
    add_bullets(
        doc,
        [
            "Complete ablation study and measured end-to-end latency (mean/P95/P99).",
            "Production deployment with horizontal scaling.",
            "Real federated deployment across independent institutions.",
            "NPCI sandbox UPI integration and regulatory compliance audit.",
        ],
    )

    # --- XII. CONCLUSION ---
    add_heading(doc, "XII. CONCLUSION", 1)
    conclusion = (
        "This paper presented SafePay, a privacy-oriented multi-layer framework for real-time "
        "financial fraud detection integrating behavioural biometrics, XGBoost ML scoring, SHAP "
        "explainability, blockchain-based anonymised intelligence sharing, and federated learning. "
    )
    if xgb:
        conclusion += (
            f"On IEEE-CIS, XGBoost achieved {fmt_metric(xgb, 'accuracy')} accuracy, "
            f"{fmt_metric(xgb, 'f1')} F1, and {fmt_metric(xgb, 'roc_auc')} ROC-AUC. "
        )
    else:
        conclusion += "On IEEE-CIS, XGBoost reached 97.55% accuracy (full metrics pending). "
    conclusion += (
        "Federated simulation achieved AUC 0.846. End-to-end platform validation confirmed "
        "auth, payment, and security controls. Remaining work: ablation evidence, latency "
        "benchmarks, and expanded primary literature before IEEE submission."
    )
    add_para(doc, conclusion)

    # --- REFERENCES ---
    add_heading(doc, "REFERENCES", 1)
    refs = [
        '[1] H. B. McMahan et al., "Communication-Efficient Learning of Deep Networks from Decentralized Data," Proc. AISTATS, 2017.',
        '[2] S. M. Lundberg and S. I. Lee, "A Unified Approach to Interpreting Model Predictions," Proc. NeurIPS, 2017.',
        '[3] N. V. Chawla et al., "SMOTE: Synthetic Minority Over-sampling Technique," J. Artif. Intell. Res., vol. 16, pp. 321–357, 2002.',
        '[4] T. Chen and C. Guestrin, "XGBoost: A Scalable Tree Boosting System," Proc. KDD, 2016.',
        '[5] IEEE-CIS Fraud Detection Dataset, Kaggle, 2019. [Online]. Available: https://www.kaggle.com/c/ieee-fraud-detection',
        '[6] A. Abdallah et al., "Financial Fraud Detection Using Machine Learning and SMOTE," IEEE Access, 2026.',
        '[7] M. Mondal and P. Bours, "Continuous Authentication Using Keystroke Dynamics," IEEE Trans. Dependable Secure Comput., 2023.',
        '[8] Y. Zheng et al., "Blockchain-Based Fraud Detection for Cross-Bank Collaboration," Financial Innovation, vol. 9, 2023.',
        '[9] P. Fiore et al., "Using Generative Adversarial Networks for Fraud Detection in IEEE-CIS," ACM ICAIF, 2019.',
        '[10] R. Jablaoui et al., "Deep Learning Enabled Intrusion Detection for IoT Security," Wireless Networks, Springer, 2025.',
        '[11] A. Buczak and E. Guven, "A Survey of Data Mining and ML Methods for Cyber Security Intrusion Detection," IEEE Commun. Surveys Tuts., 2016.',
        '[12] Y. Wu and B. Zou, "Deep Learning-Based Intrusion Detection: Status and Trends," J. Imaging, vol. 10, no. 10, 2024.',
        '[13] N. Latif and W. Ma, "Securing Federated Learning: Advances and Challenges," Artif. Intell. Rev., 2025.',
        '[14] Reserve Bank of India, Annual Report 2024 — Digital Payments and Fraud Statistics.',
        '[15] M. Beutel et al., "Flower: A Friendly Federated Learning Framework," arXiv:2007.14390, 2020.',
        '[16] G. Ke et al., "LightGBM: A Highly Efficient Gradient Boosting Decision Tree," Proc. NeurIPS, 2017.',
        '[17] F. Pedregosa et al., "Scikit-learn: Machine Learning in Python," JMLR, vol. 12, pp. 2825–2830, 2011.',
        '[18] S. Ghosh et al., "Explainable AI for Financial Fraud Detection: A Systematic Review," Expert Syst. Appl., 2024.',
        '[19] E. C. Pinto Neto et al., "Deep Learning for Intrusion Detection in Emerging Technologies: A Survey," Artif. Intell. Rev., 2025.',
        '[20] FastAPI Documentation. [Online]. Available: https://fastapi.tiangolo.com',
        '[21] Hardhat Ethereum Development Environment. [Online]. Available: https://hardhat.org',
        '[22] SHAP Documentation. [Online]. Available: https://shap.readthedocs.io',
    ]
    for ref in refs:
        add_para(doc, ref)

    try:
        doc.save(OUTPUT)
        saved = OUTPUT
    except PermissionError:
        doc.save(OUTPUT_FALLBACK)
        saved = OUTPUT_FALLBACK
        print(f"WARNING: Could not overwrite {OUTPUT.name} (file may be open in Word).")
    print(f"Saved: {saved}")
    if metrics_data:
        print("Included ML metrics from ml_evaluation_results.json")
    else:
        print("ML metrics not found — Tables II/III show TBD. Run evaluate_models.py first.")


if __name__ == "__main__":
    build_document()
